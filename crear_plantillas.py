"""
crear_plantillas.py
Genera las plantillas Word estandarizadas con marcadores {{CAMPO}}
a partir de los documentos originales del área de Talento EAFIT.
"""
import re
import copy
from docx import Document
from docx.shared import Pt
import os

UPLOADS = "/sessions/awesome-modest-lovelace/mnt/uploads"
TEMPLATES = os.path.join(os.path.dirname(__file__), "templates")
os.makedirs(TEMPLATES, exist_ok=True)

# ─── Mapa de reemplazos: patrón regex → marcador ────────────────────────────
REEMPLAZOS_COMUNES = [
    # Representante legal EAFIT (datos fijos, no se reemplazan)
    # Nombre empresa / escenario
    (r'\b_____+\b(?=.*ciudad|.*direc)', '{{CIUDAD_EMPRESA}}'),  # intento contextual

    # Guiones de la cláusula PRIMERA (ubicación)
    (r'ubicado en ___+', 'ubicado en {{CIUDAD_EMPRESA}}'),
    # Fechas de duración
    (r'entre el ____? de _____+ de 20__', 'entre el {{DIA_INICIO}} de {{MES_INICIO}} de {{ANIO_INICIO}}'),
    (r'y el ____? de ______+ de 20___', 'y el {{DIA_FIN}} de {{MES_FIN}} de {{ANIO_FIN}}'),
    # Monto auxilio
    (r'la suma mensual de _______+ PESOS \(\$_______\)', 
     'la suma mensual de {{MONTO_LETRAS}} PESOS (${{MONTO_NUMEROS}})'),
    # Tutor empresa
    (r'designa como TUTOR a ___+,\s*identificado\(a\) con la cédula de ciudadanía No\. __+',
     'designa como TUTOR a {{NOMBRE_TUTOR}}, identificado(a) con la cédula de ciudadanía No. {{CEDULA_TUTOR}}'),
    # Monitor EAFIT
    (r'designa como MONITOR a ___+,?\s*identificad[ao] con la cédula de ciudadanía No\.__+',
     'designa como MONITOR a {{NOMBRE_MONITOR}}, identificada con la cédula de ciudadanía No. {{CEDULA_MONITOR}}'),
    # Firma final
    (r'suscriben este documento el ___+\s*\(___\) de ___+ de _____\.',
     'suscriben este documento el {{DIA_FIRMA}} ({{DIA_FIRMA_LETRAS}}) de {{MES_FIRMA}} de {{ANIO_FIRMA}}.'),
    # Actividades (guiones simples)
    (r'^_+\.$', '{{ACTIVIDAD}}'),
]

MARCADORES_ENCABEZADO = {
    # En la descripción de partes (párrafo 2)
    'PRIMERA. OBJETO': None,  # marcador para ubicación
}


def reemplazar_en_parrafo(para, doc_tipo):
    """Reemplaza blancos en un párrafo con marcadores."""
    texto = para.text
    nuevo = texto

    # Ubicación empresa (cláusula PRIMERA)
    nuevo = re.sub(r'ubicado en _+', 'ubicado en {{CIUDAD_EMPRESA}}', nuevo)
    
    # Fechas inicio/fin
    nuevo = re.sub(
        r'entre el _+ de _+ de 20_+',
        'entre el {{DIA_INICIO}} de {{MES_INICIO}} de {{ANIO_INICIO}}', nuevo)
    nuevo = re.sub(
        r'y el _+ de _+ de 20_+',
        'y el {{DIA_FIN}} de {{MES_FIN}} de {{ANIO_FIN}}', nuevo)
    
    # Monto (solo remunerados)
    nuevo = re.sub(
        r'la suma mensual de _+ PESOS \(\$_+\)',
        'la suma mensual de {{MONTO_LETRAS}} PESOS (${{MONTO_NUMEROS}})', nuevo)
    
    # Tutor empresa
    nuevo = re.sub(
        r'TUTOR a _+,?\s*identificado\(a\) con la cédula de ciudadanía No\.\s*_+',
        'TUTOR a {{NOMBRE_TUTOR}}, identificado(a) con la cédula de ciudadanía No. {{CEDULA_TUTOR}}', nuevo)
    
    # Monitor EAFIT
    nuevo = re.sub(
        r'MONITOR a _+,?\s*identificad[ao] con la cédula de ciudadanía No\._+\s*\.',
        'MONITOR a {{NOMBRE_MONITOR}}, identificada con la cédula de ciudadanía No. {{CEDULA_MONITOR}} .', nuevo)
    
    # Firma final
    nuevo = re.sub(
        r'suscriben este documento el _+\s*\(_+\) de _+ de _+\.',
        'suscriben este documento el {{DIA_FIRMA}} ({{DIA_FIRMA_LETRAS}}) de {{MES_FIRMA}} de {{ANIO_FIRMA}}.', nuevo)
    
    # Actividades (líneas solo con guiones)
    if re.match(r'^_+\.$', nuevo.strip()):
        nuevo = '{{ACTIVIDAD}}'
    
    # Actualizar el run si hubo cambio
    if nuevo != texto and para.runs:
        # Poner todo en el primer run, limpiar el resto
        para.runs[0].text = nuevo
        for run in para.runs[1:]:
            run.text = ''
    elif nuevo != texto:
        para.clear()
        para.add_run(nuevo)


def crear_plantilla(src_filename, dest_filename, doc_tipo):
    src = os.path.join(UPLOADS, src_filename)
    dest = os.path.join(TEMPLATES, dest_filename)
    doc = Document(src)
    
    # Reemplazar en encabezado: párrafo 2 tiene datos de EAFIT (fijos) + datos empresa (variable)
    # Párrafo 2: "Entre los suscritos a saber, ISABEL... y {NOMBRE_EMPRESA}, {TIPO_SOCIEDAD}..."
    partes_para = doc.paragraphs[2]
    texto_partes = partes_para.text
    # El texto tiene datos EAFIT fijos hasta "denominada" o "y" conectando con la empresa
    # Buscamos el patrón después del representante de EAFIT
    nuevo_partes = re.sub(
        r'(facultada para actuar en nombre.*?;)\s+y\s+(.*?)(?=,\s+representad)',
        r'\1 y {{NOMBRE_EMPRESA}}, {{TIPO_SOCIEDAD}}, identificada con NIT No. {{NIT_EMPRESA}}',
        texto_partes, flags=re.DOTALL
    )
    # Representante legal empresa
    nuevo_partes = re.sub(
        r'representad[ao] por\s+([A-ZÁÉÍÓÚÑ\s]+),\s+identificad[ao] con cédula de ciudadanía No\.?\s*[\d.,]+',
        'representada por {{NOMBRE_REPRESENTANTE}}, identificado(a) con cédula de ciudadanía No. {{CEDULA_REPRESENTANTE}}',
        nuevo_partes
    )
    # Cargo representante
    nuevo_partes = re.sub(
        r'en su calidad de\s+([A-Za-záéíóúÁÉÍÓÚñÑ\s]+),?\s+quien declara estar',
        'en su calidad de {{CARGO_REPRESENTANTE}}, quien declara estar',
        nuevo_partes
    )
    # Datos estudiante
    nuevo_partes = re.sub(
        r'EL ESTUDIANTE:\s+([A-ZÁÉÍÓÚÑ\s]+),?\s+identificad[ao] con cédula de ciudadanía No\.?\s*[\d.,]+',
        'EL ESTUDIANTE: {{NOMBRE_ESTUDIANTE}}, identificado(a) con cédula de ciudadanía No. {{CEDULA_ESTUDIANTE}}',
        nuevo_partes
    )
    nuevo_partes = re.sub(
        r'estudiante del programa\s+([^\.,]+)',
        'estudiante del programa {{PROGRAMA_ACADEMICO}}',
        nuevo_partes
    )
    
    if nuevo_partes != texto_partes and partes_para.runs:
        partes_para.runs[0].text = nuevo_partes
        for r in partes_para.runs[1:]:
            r.text = ''
    
    # Reemplazar en todos los párrafos
    for para in doc.paragraphs:
        reemplazar_en_parrafo(para, doc_tipo)
    
    # También en tablas si las hay
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    reemplazar_en_parrafo(para, doc_tipo)
    
    doc.save(dest)
    print(f"✓ Creada: {dest_filename}")


# Crear las 4 plantillas de convenio de práctica
crear_plantilla(
    '01. CONVENIO DE VINCULACIÓN FORMATIVA REMUNERADA Y ORGANIZACIÓN paga ARL.docx',
    'convenio_01_remunerado_org_arl.docx',
    'remunerado_org'
)
crear_plantilla(
    '02. CONVENIO DE VINCULACIÓN FORMATIVA REMUNERADA cuando EAFIT paga ARL actualizada.DOCX',
    'convenio_02_remunerado_eafit_arl.docx',
    'remunerado_eafit'
)
crear_plantilla(
    '03. CONVENIO DE VINCULACIÓN FORMATIVA NO REMUNERADA Y ORGANIZACIÓN paga ARL actualizada.DOCX',
    'convenio_03_no_remunerado_org_arl.docx',
    'no_remunerado_org'
)
crear_plantilla(
    '04. CONVENIO DE VINCULACIÓN FORMATIVA NO REMUNERADA Y EAFIT paga ARL actualizada.DOCX',
    'convenio_04_no_remunerado_eafit_arl.docx',
    'no_remunerado_eafit'
)

print("\n✅ Todas las plantillas creadas exitosamente.")
