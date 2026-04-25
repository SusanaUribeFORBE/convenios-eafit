"""
generador.py — Motor de generacion de convenios EAFIT
"""
import re, os, io
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE      = os.path.dirname(os.path.abspath(__file__))
TEMPLATES = BASE

MESES = {1:'enero',2:'febrero',3:'marzo',4:'abril',5:'mayo',6:'junio',
         7:'julio',8:'agosto',9:'septiembre',10:'octubre',11:'noviembre',12:'diciembre'}

NUMEROS = {1:'UN',2:'DOS',3:'TRES',4:'CUATRO',5:'CINCO',6:'SEIS',7:'SIETE',
           8:'OCHO',9:'NUEVE',10:'DIEZ',11:'ONCE',12:'DOCE',13:'TRECE',
           14:'CATORCE',15:'QUINCE',16:'DIECISEIS',17:'DIECISIETE',
           18:'DIECIOCHO',19:'DIECINUEVE',20:'VEINTE',21:'VEINTIUN',
           22:'VEINTIDOS',23:'VEINTITRES',24:'VEINTICUATRO',25:'VEINTICINCO',
           26:'VEINTISEIS',27:'VEINTISIETE',28:'VEINTIOCHO',29:'VEINTINUEVE',
           30:'TREINTA',31:'TREINTA Y UNO'}


def seleccionar_plantilla(tipo, remunerada, arl):
    if tipo == "Practica":
        if remunerada:
            return "convenio_01_remunerado_org_arl.docx" if arl == "Organizacion" else "convenio_02_remunerado_eafit_arl.docx"
        else:
            return "convenio_03_no_remunerado_org_arl.docx" if arl == "Organizacion" else "convenio_04_no_remunerado_eafit_arl.docx"
    return "convenio_01_remunerado_org_arl.docx"


def eliminar_comentarios(doc):
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for tag in [f'{{{W}}}commentRangeStart', f'{{{W}}}commentRangeEnd', f'{{{W}}}commentReference']:
        for elem in list(doc.element.body.iter(tag)):
            p = elem.getparent()
            if p is not None:
                p.remove(elem)
    for run in list(doc.element.body.iter(f'{{{W}}}r')):
        if not [h for h in run if not h.tag.endswith('}rPr')]:
            p = run.getparent()
            if p is not None:
                p.remove(run)


def reemplazar_parrafo(para, reemplazos):
    texto = para.text
    nuevo = texto
    for k, v in reemplazos.items():
        nuevo = nuevo.replace(k, v)
    if nuevo != texto and para.runs:
        para.runs[0].text = nuevo
        for r in para.runs[1:]:
            r.text = ''


def generar_documento(datos):
    # Normalizar tipo y ARL (quitar tildes para comparar)
    tipo = datos.get('tipo_experiencia', 'Practica').replace('a\u0301','a').replace('c','c')
    rem  = datos.get('remunerada', False)
    arl  = datos.get('quien_paga_arl', 'Organizacion').replace('o\u0301','o').replace('n','n')

    plantilla = seleccionar_plantilla(tipo, rem, arl)
    doc = Document(os.path.join(TEMPLATES, plantilla))

    fi     = datos['fecha_inicio']
    ff     = datos['fecha_fin']
    ffirma = datos.get('fecha_firma', ff) or ff

    reemplazos = {
        '{{NOMBRE_EMPRESA}}':       datos.get('nombre_empresa',''),
        '{{TIPO_SOCIEDAD}}':        datos.get('tipo_sociedad','S.A.S.'),
        '{{NIT_EMPRESA}}':          datos.get('nit_empresa',''),
        '{{NOMBRE_REPRESENTANTE}}': datos.get('nombre_representante',''),
        '{{CEDULA_REPRESENTANTE}}': datos.get('cedula_representante',''),
        '{{CARGO_REPRESENTANTE}}':  datos.get('cargo_representante',''),
        '{{CIUDAD_EMPRESA}}':       datos.get('ciudad_empresa',''),
        '{{NOMBRE_ESTUDIANTE}}':    datos.get('nombre_estudiante',''),
        '{{CEDULA_ESTUDIANTE}}':    datos.get('cedula_estudiante',''),
        '{{PROGRAMA_ACADEMICO}}':   datos.get('programa_academico',''),
        '{{DIA_INICIO}}':           str(fi.day),
        '{{MES_INICIO}}':           MESES[fi.month],
        '{{ANIO_INICIO}}':          str(fi.year),
        '{{DIA_FIN}}':              str(ff.day),
        '{{MES_FIN}}':              MESES[ff.month],
        '{{ANIO_FIN}}':             str(ff.year),
        '{{NOMBRE_TUTOR}}':         datos.get('nombre_tutor',''),
        '{{CEDULA_TUTOR}}':         datos.get('cedula_tutor',''),
        '{{NOMBRE_MONITOR}}':       datos.get('nombre_monitor',''),
        '{{CEDULA_MONITOR}}':       datos.get('cedula_monitor',''),
        '{{DIA_FIRMA}}':            str(ffirma.day),
        '{{DIA_FIRMA_LETRAS}}':     NUMEROS.get(ffirma.day, str(ffirma.day)),
        '{{MES_FIRMA}}':            MESES[ffirma.month],
        '{{ANIO_FIRMA}}':           str(ffirma.year),
        '{{MONTO_LETRAS}}':         datos.get('monto_letras','').upper(),
        '{{MONTO_NUMEROS}}':        datos.get('monto_numeros',''),
    }

    actividades = [a for a in datos.get('actividades',[]) if a.strip()]

    # Las plantillas tienen 8 marcadores {{ACTIVIDAD}} (maximo permitido).
    # Reemplazar los que tengan actividad; eliminar los sobrantes.
    paras_actividad = [p for p in doc.paragraphs if '{{ACTIVIDAD}}' in p.text]

    for i, para in enumerate(paras_actividad):
        if i < len(actividades):
            if para.runs:
                para.runs[0].text = para.text.replace('{{ACTIVIDAD}}', actividades[i])
                for r in para.runs[1:]:
                    r.text = ''
        else:
            para._element.getparent().remove(para._element)

    for para in doc.paragraphs:
        reemplazar_parrafo(para, reemplazos)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    reemplazar_parrafo(para, reemplazos)

    eliminar_comentarios(doc)

    # Limpieza final: eliminar cualquier parrafo numerado vacio que haya quedado
    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for p in list(doc.element.body.iter(f'{{{W_NS}}}p')):
        tiene_texto = any(
            (t.text or '').strip()
            for t in p.iter(f'{{{W_NS}}}t')
        )
        tiene_num = p.find(f'.//{{{W_NS}}}numPr') is not None
        if not tiene_texto and tiene_num:
            padre = p.getparent()
            if padre is not None:
                padre.remove(p)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
